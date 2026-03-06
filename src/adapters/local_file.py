"""Local file adapter for PDF, Word, text files."""
import logging
import os
from .base import BaseAdapter, SourceContent

logger = logging.getLogger(__name__)


class LocalFileAdapter(BaseAdapter):
    name = "local_file"

    _EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".rst", ".csv", ".xlsx"}

    @staticmethod
    def _normalize_path(source: str) -> str:
        """Normalize file path: expand ~ and resolve relative paths."""
        return os.path.abspath(os.path.expanduser(source))

    def can_handle(self, source: str) -> bool:
        if source.startswith(("http://", "https://", "search:")):
            return False
        _, ext = os.path.splitext(source.lower())
        if ext in self._EXTENSIONS:
            return True
        normalized = self._normalize_path(source)
        return os.path.isfile(normalized)

    async def extract(self, source: str) -> SourceContent:
        source = self._normalize_path(source)
        if not os.path.exists(source):
            raise FileNotFoundError(f"File not found: {source}")

        _, ext = os.path.splitext(source.lower())
        title = os.path.basename(source)

        if ext == ".pdf":
            text = self._read_pdf(source)
        elif ext in (".docx", ".doc"):
            text = self._read_docx(source)
        elif ext in (".xlsx", ".csv"):
            text = self._read_table(source, ext)
        else:
            with open(source, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        return SourceContent(
            text=text,
            title=title,
            source_url=f"file://{os.path.abspath(source)}",
            source_type="local_file",
        )

    def _read_pdf(self, path: str) -> str:
        try:
            import pymupdf
            doc = pymupdf.open(path)
            return "\n".join(page.get_text() for page in doc)
        except ImportError:
            try:
                from pypdf import PdfReader
                reader = PdfReader(path)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                raise ImportError("pip install pymupdf or pypdf")

    def _read_docx(self, path: str) -> str:
        _, ext = os.path.splitext(path.lower())

        # Try mammoth first (best quality, handles both formatting and tables)
        try:
            import mammoth
            with open(path, "rb") as f:
                result = mammoth.convert_to_markdown(f)
                text = result.value
                if result.messages:
                    for msg in result.messages:
                        logger.debug("mammoth: %s", msg)
                if text and text.strip():
                    return text
        except ImportError:
            pass
        except Exception as e:
            logger.warning("mammoth failed for %s: %s", path, e)

        # Fallback: python-docx (paragraphs + tables)
        try:
            from docx import Document
            doc = Document(path)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))
                    parts.append("")
            text = "\n".join(parts)
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception as e:
            logger.warning("python-docx failed for %s: %s", path, e)

        # Fallback for .doc (old format): try antiword or textract
        if ext == ".doc":
            try:
                import subprocess
                result = subprocess.run(
                    ["antiword", path], capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

            try:
                import subprocess
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "txt:Text", "--outdir", "/tmp", path],
                    capture_output=True, text=True, timeout=60
                )
                txt_path = os.path.join("/tmp", os.path.splitext(os.path.basename(path))[0] + ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    os.remove(txt_path)
                    if text.strip():
                        return text
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

        raise ImportError(
            f"Word 파일을 읽을 수 없습니다. 다음 중 하나를 설치하세요:\n"
            f"  pip install mammoth    (권장, .docx)\n"
            f"  pip install python-docx (.docx)\n"
            f"  apt install antiword   (.doc 구형 포맷)"
        )

    def _read_table(self, path: str, ext: str) -> str:
        try:
            import pandas as pd
            if ext == ".csv":
                df = pd.read_csv(path)
            else:
                df = pd.read_excel(path)
            return df.to_markdown(index=False)
        except ImportError:
            raise ImportError("pip install pandas openpyxl")
